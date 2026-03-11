# -*- coding: utf-8 -*-
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

IMG_DIR = r"d:\desktop\picture prompt\mv_images"
OUTPUT_PATH = r"d:\desktop\picture prompt\MV分镜大纲_最终版.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading_style(doc, text, level=1, color='1a1a2e'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(13)
    r, g, b = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    return p

def add_section_divider(doc, text, bg_color='2d1b4e', text_color='ffffff'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    r, g, b = tuple(int(text_color[i:i+2], 16) for i in (0, 2, 4))
    run.font.color.rgb = RGBColor(r, g, b)
    doc.add_paragraph()

def add_shot(doc, shot_num, time_range, img_path, shot_type, description, video_prompt, duration, bg_color='f5f0ff'):
    """添加单个分镜"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 左列放图片
    left_cell = table.cell(0, 0)
    left_cell.width = Inches(2.5)
    set_cell_bg(left_cell, 'e8e0f0')
    
    if img_path and os.path.exists(img_path):
        p_img = left_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(2.3))
    else:
        p_img = left_cell.paragraphs[0]
        p_img.add_run(f"[图片{shot_num}]")
    
    # 右列放文字内容
    right_cell = table.cell(0, 1)
    set_cell_bg(right_cell, bg_color)
    
    # 分镜号 + 时间
    p_header = right_cell.add_paragraph()
    r1 = p_header.add_run(f"【分镜 {shot_num:02d}】 ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x2d, 0x1b, 0x4e)
    r2 = p_header.add_run(f"{time_range}  [{duration}]")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x88, 0x66, 0xaa)
    
    # 镜头类型
    p_type = right_cell.add_paragraph()
    r_label = p_type.add_run("镜头：")
    r_label.bold = True
    r_label.font.size = Pt(10)
    r_label.font.color.rgb = RGBColor(0x55, 0x33, 0x88)
    r_content = p_type.add_run(shot_type)
    r_content.font.size = Pt(10)
    r_content.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 画面描述
    p_desc = right_cell.add_paragraph()
    r_label2 = p_desc.add_run("画面：")
    r_label2.bold = True
    r_label2.font.size = Pt(10)
    r_label2.font.color.rgb = RGBColor(0x55, 0x33, 0x88)
    r_desc = p_desc.add_run(description)
    r_desc.font.size = Pt(10)
    r_desc.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # 视频提示词
    p_prompt = right_cell.add_paragraph()
    r_label3 = p_prompt.add_run("▶ Seedance提示词：")
    r_label3.bold = True
    r_label3.font.size = Pt(9)
    r_label3.font.color.rgb = RGBColor(0x0a, 0x7c, 0x5a)
    
    p_prompt_content = right_cell.add_paragraph()
    r_prompt = p_prompt_content.add_run(video_prompt)
    r_prompt.font.size = Pt(9)
    r_prompt.italic = True
    r_prompt.font.color.rgb = RGBColor(0x1a, 0x5c, 0x3a)
    
    # 删除右侧单元格第一个空段落
    first_p = right_cell.paragraphs[0]
    if not first_p.text:
        p = first_p._element
        p.getparent().remove(p)
    
    doc.add_paragraph()

def add_group_shot(doc, group_num, time_range, img_paths, shot_type, description, video_prompt, total_duration, img_names=None):
    """添加组合分镜（多图横排+长提示词）"""
    # 标题行
    table_header = doc.add_table(rows=1, cols=1)
    table_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcell = table_header.cell(0, 0)
    set_cell_bg(hcell, '4a2080')
    p_h = hcell.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_h = p_h.add_run(f"  ▶ 组合片段 {group_num}  {time_range}  [{total_duration}]  {shot_type}")
    r_h.bold = True
    r_h.font.size = Pt(11)
    r_h.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    
    # 图片横排
    valid_paths = [p for p in img_paths if p and os.path.exists(p)]
    if valid_paths:
        img_table = doc.add_table(rows=1, cols=len(valid_paths))
        img_table.style = 'Table Grid'
        img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, ip in enumerate(valid_paths):
            cell = img_table.cell(0, i)
            set_cell_bg(cell, 'd8ccea')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            width = min(Inches(1.6), Inches(7.0 / len(valid_paths)))
            run.add_picture(ip, width=width)
            if img_names and i < len(img_names):
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.add_run(img_names[i])
                r2.font.size = Pt(8)
                r2.font.color.rgb = RGBColor(0x66, 0x44, 0x88)
    
    # 描述和提示词
    desc_table = doc.add_table(rows=1, cols=1)
    desc_table.style = 'Table Grid'
    desc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    dcell = desc_table.cell(0, 0)
    set_cell_bg(dcell, 'ede8f8')
    
    p_desc = dcell.add_paragraph()
    r_dl = p_desc.add_run("画面描述：")
    r_dl.bold = True
    r_dl.font.size = Pt(10)
    r_dl.font.color.rgb = RGBColor(0x55, 0x33, 0x88)
    r_dc = p_desc.add_run(description)
    r_dc.font.size = Pt(10)
    r_dc.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    p_prompt_label = dcell.add_paragraph()
    r_pl = p_prompt_label.add_run("▶ Seedance提示词（组合生成）：")
    r_pl.bold = True
    r_pl.font.size = Pt(9)
    r_pl.font.color.rgb = RGBColor(0x0a, 0x7c, 0x5a)
    
    p_prompt_val = dcell.add_paragraph()
    r_pv = p_prompt_val.add_run(video_prompt)
    r_pv.font.size = Pt(9)
    r_pv.italic = True
    r_pv.font.color.rgb = RGBColor(0x1a, 0x5c, 0x3a)
    
    # 删第一个空段落
    first_p = dcell.paragraphs[0]
    if not first_p.text:
        p = first_p._element
        p.getparent().remove(p)
    
    doc.add_paragraph()

def img(n):
    return os.path.join(IMG_DIR, f"img_{n:02d}.png")

# ============================================================
# 开始创建文档
# ============================================================
doc = Document()

# 设置页边距
from docx.shared import Mm
sections = doc.sections
for section in sections:
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

# ========== 封面 ==========
doc.add_paragraph()
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("青春如烟，呼吸是蓝")
r_title.bold = True
r_title.font.size = Pt(28)
r_title.font.color.rgb = RGBColor(0x2d, 0x1b, 0x4e)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_subtitle.add_run("MV 分镜大纲 · 50秒完整版")
r_sub.font.size = Pt(16)
r_sub.font.color.rgb = RGBColor(0x66, 0x44, 0x88)

p_bgm = doc.add_paragraph()
p_bgm.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_bgm = p_bgm.add_run("BGM：Phase to Phrase — TK from 凛として時雨")
r_bgm.font.size = Pt(12)
r_bgm.font.color.rgb = RGBColor(0x88, 0x66, 0xaa)

doc.add_paragraph()

# 风格说明
style_table = doc.add_table(rows=1, cols=1)
style_cell = style_table.cell(0, 0)
set_cell_bg(style_cell, '1a0a30')
p_style = style_cell.paragraphs[0]
p_style.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_style = p_style.add_run(
    "风格基调：王家卫×岩井俊二×莉莉周  /  胶片质感·色彩浓郁·意识流闪切\n"
    "节奏结构：开头静谧听歌 → 突然爆发快节奏 → 情绪碎片交替 → 尾声消散"
)
r_style.font.size = Pt(11)
r_style.font.color.rgb = RGBColor(0xdd, 0xcc, 0xff)
doc.add_paragraph()

# ========== 节奏总谱 ==========
add_heading_style(doc, "50秒节奏总谱", level=2, color='2d1b4e')

timeline_data = [
    ("段落", "时间", "情绪", "节奏", "镜头风格"),
    ("INTRO 开场", "0-8s", "静谧·孤独", "慢 ●○○", "长镜头·空旷感"),
    ("BREAK 爆发", "8-10s", "撕裂·转折", "冲击 ●●●", "闪切·过曝"),
    ("VERSE 1 青春碎片", "10-25s", "心跳·疯跑·爱恋", "快 ●●○", "手持抖动·快剪"),
    ("VERSE 2 伤痛记忆", "25-38s", "孤独·考试·压抑", "中快 ●●○", "推镜·特写"),
    ("CHORUS 情绪高潮", "38-46s", "爆发·毕业·燃烧", "极快 ●●●", "大范围运镜"),
    ("OUTRO 消散", "46-50s", "余韵·漂浮", "慢 ●○○", "叠化·渐出"),
]

tl_table = doc.add_table(rows=len(timeline_data), cols=5)
tl_table.style = 'Table Grid'
tl_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for row_idx, row_data in enumerate(timeline_data):
    for col_idx, cell_text in enumerate(row_data):
        cell = tl_table.cell(row_idx, col_idx)
        if row_idx == 0:
            set_cell_bg(cell, '3d0a6e')
        elif row_idx % 2 == 0:
            set_cell_bg(cell, 'ede8f8')
        else:
            set_cell_bg(cell, 'f8f4ff')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        if row_idx == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.size = Pt(9)

doc.add_paragraph()

# ========== INTRO 段落 ==========
add_section_divider(doc, "═══  INTRO 开场  |  0s – 8s  |  静谧·孤独·等待  ═══", bg_color='1a0a30', text_color='d4b8ff')

add_shot(doc, 1, "0:00 – 0:05", img(40), "面部特写·慢推",
    "少年戴着旧式大耳机，双眼微闭，蓝色天空背景，发丝随风微动，光斑渗入镜头边缘。开场静止感，只有音乐在流动。",
    "@图片1 作为首帧，胶片粒子质感，少年戴耳机侧颜特写，蓝色天空漫溢过曝，发丝随风轻颤，镜头极缓慢Slow push in推向面部，光晕在镜头角落慢慢扩散，胶片刮痕隐约闪现，画面几乎静止，背景虚化，只留呼吸感。5秒，16:9，无台词，环境音：远处风声+耳机微弱漏音。禁止：任何文字字幕",
    "5s", 'f0e8ff')

add_shot(doc, 2, "0:05 – 0:08", img(39), "中景·固定镜头",
    "火车车厢内，少年少女并排坐，各自戴耳机，阳光从侧窗炸裂射入，两人头部在逆光中成黑色剪影，座椅的红色形成对比。",
    "@图片1 作为首帧，胶片质感，火车车厢内两人并排坐听歌，侧窗射入强烈逆光，人物轮廓半剪影，红色座椅与蓝绿色调形成对比，镜头固定，轻微手持抖动，窗外景物模糊流动，气氛安静克制，阳光光晕在玻璃上渗开。3秒，16:9。禁止：任何文字字幕",
    "3s", 'f0e8ff')

# ========== BREAK 爆发转折 ==========
add_section_divider(doc, "═══  BREAK 爆发  |  8s – 10s  |  音乐突然炸开  ═══", bg_color='4a0080', text_color='ffccff')

add_shot(doc, 3, "0:08 – 0:10", img(28), "极速闪切·过曝",
    "少年在户外逆光中张嘴大喊，强光爆炸性覆盖镜头，过曝瞬间把画面吞没。音乐在此炸开，节奏猛地切换。",
    "@图片1 作为首帧，少年在户外逆光中仰头大喊，太阳从脑后爆出强光将画面过曝为白，闪切2次（每次0.5秒），镜头轻微Handheld抖动，色调饱和度骤增，画面边缘出现胶片烧痕效果。2秒，16:9，音效：音乐突然爆发。禁止：任何文字字幕",
    "2s", 'e8d0ff')

# ========== VERSE 1 段落 ==========
add_section_divider(doc, "═══  VERSE 1 青春碎片  |  10s – 25s  |  心跳·疯跑·爱恋  ═══", bg_color='2d1b4e', text_color='ffddee')

add_shot(doc, 4, "0:10 – 0:15", img(30), "低角度仰拍·跟随",
    "少年背着书包在铁路旁全力奔跑，侧面低角度，火车从背后呼啸而过，动作模糊带有速度感，蓝天在画面顶部。",
    "@图片1 作为首帧，低角度侧面跟拍少年奔跑，铁路旁火车从背景高速掠过，镜头低角度仰拍Tracking跟随，运动模糊强，胶片颗粒，天空蓝色与铁路黄色形成对比，速度感强烈，脚步声与火车轰鸣叠加。5秒，16:9。禁止：任何文字字幕",
    "5s", 'fff0f5')

# 组合分镜：稻田玩耍组
add_group_shot(doc, 1, "0:15 – 0:22", [img(1), img(6), img(7), img(22)],
    "闪切·快剪·稻田系列",
    "稻田中一组快切：少年戴耳机听歌(img01)→稻田中奔跑(img06)→花田里(img07)→稻田背影(img22)，每个镜头约1.5-2秒，形成青春活力的碎片感。",
    """15秒快剪青春碎片，胶片粒子+高饱和，全程Handheld手持晃动，
@图片1 少年在稻田戴耳机听歌，面部特写慢推，2秒；
@图片2 少年在稻田中奔跑，低角度跟随，1.5秒，脚踢稻叶慢镜；
@图片3 花田全景，少年穿过花丛，镜头从花丛里仰拍，1.5秒；
@图片4 稻田背影远景，Pan right横摇，2秒，色调翠绿过曝。
全程快切，每次切换用帧闪白，音乐节奏卡点。禁止：任何文字字幕""",
    "7s", [f"img_{n:02d}" for n in [1, 6, 7, 22]])

add_shot(doc, 5, "0:22 – 0:25", img(11), "手部特写·慢镜",
    "两人手指相握的瞬间特写，光斑在皮肤上晃动，胶片色调。心动的刹那定格。",
    "@图片1 作为首帧，两人手指相触特写，暖光在皮肤上散开形成光晕，慢镜头0.5倍速，轻微Zoom in，胶片色调温暖溢出，背景完全虚化，只剩手部细节。3秒，16:9，音效：心跳声渐入。禁止：任何文字字幕",
    "3s", 'fff5ee')

# ========== VERSE 2 段落 ==========
add_section_divider(doc, "═══  VERSE 2 伤痛记忆  |  25s – 38s  |  孤独·高考·压抑  ═══", bg_color='0a1a3e', text_color='aaccff')

add_shot(doc, 6, "0:25 – 0:29", img(27), "固定·走廊",
    "少年穿白色校服独自坐在昏暗走廊地上，窗外透来蓝绿色冷光，膝盖抱胸，眼神空洞向下。",
    "@图片1 作为首帧，胶片粒子青绿色调，少年独坐学校走廊地面，膝盖收缩，光从两侧窗户透入形成对称光带，镜头固定中景，轻微Slow push in推向面部，地板反光，周围空旷寂静，色调冷绿压抑。4秒，16:9，音效：空旷走廊回声。禁止：任何文字字幕",
    "4s", 'e8f0ff')

add_shot(doc, 7, "0:29 – 0:33", img(29), "俯拍·考试",
    "考场内少年坐在试卷前发呆，鱼眼镜头微微畸变，周围同学埋头写字，他的视线空洞，窗外强光。",
    "@图片1 作为首帧，高考考场，鱼眼广角轻微畸变，少年在试卷前发呆，周围同学埋头答题，窗外逆光过曝，镜头缓缓从少年面部Slow pull back拉远，同时色温逐渐变冷，压抑窒息感。4秒，16:9，音效：笔尖摩擦纸张声渐弱。禁止：任何文字字幕",
    "4s", 'e8f0ff')

# 组合分镜：打工/社会压迫组
add_group_shot(doc, 2, "0:33 – 0:38", [img(38), img(19), img(16)],
    "快切·成年压力",
    "长大后的碎片：音像店里发呆(img38)→孤坐台阶(img19)→打电话(img16)，代表走出青春进入现实的压力与孤独。",
    """5秒快切压抑碎片，冷色调胶片粒子，
@图片1 音像店走廊，少年低着头从碟片架间穿过，固定镜头，昏黄荧光灯闪烁，1.5秒；
@图片2 台阶独坐，中景固定，少年抱膝，夜灯投下长影，1.5秒；
@图片3 少年手持手机特写，手微颤，画面快速闪切2次，2秒。
切换用帧闪黑，节奏稍快，冷暖色交替营造时间跳跃感。禁止：任何文字字幕""",
    "5s", [f"img_{n:02d}" for n in [38, 19, 16]])

# ========== CHORUS 高潮段落 ==========
add_section_divider(doc, "═══  CHORUS 情绪高潮  |  38s – 46s  |  爆发·毕业·燃烧  ═══", bg_color='4a1a00', text_color='ffdd88')

add_shot(doc, 8, "0:38 – 0:41", img(37), "仰拍·反叛",
    "少年站在教室课桌上，仰拍，窗外强光从背后打来成剪影，动作张扬反叛，其他同学坐着抬头看。",
    "@图片1 作为首帧，仰拍Low angle，少年站在教室课桌上，背对强烈逆光窗户，形成剪影轮廓，光晕在身体周围爆开，镜头缓慢Crane up升起，同时环境音消失，只剩音乐，色调过曝变暖。3秒，16:9。禁止：任何文字字幕",
    "3s", 'fff8e8')

add_shot(doc, 9, "0:41 – 0:44", img(34), "大广角·航拍感",
    "高中毕业操场，纸片书本漫天飞舞，同学们欢呼，强烈的阳光把整个画面烧成白色，最欢乐最痛苦的瞬间。",
    "@图片1 作为首帧，高中操场毕业瞬间，书本纸片漫天飞舞，同学群体欢呼，广角低机位仰拍，阳光直射镜头形成强烈光晕，色调过曝偏白，镜头快速Crane up升起如航拍，胶片颗粒感强，配合音乐高潮爆发。3秒，16:9，音效：人群欢呼声混入。禁止：任何文字字幕",
    "3s", 'fff8e8')

# 组合分镜：燃烧收尾组
add_group_shot(doc, 3, "0:44 – 0:46", [img(35), img(33)],
    "快切·烟花·剪影",
    "两组对比：夜晚烟花(img35)→电线杆下两人剪影奔跑(img33)，节奏最快的瞬间，情绪最高点。",
    """2秒极速闪切，高饱和过曝，
@图片1 夜晚两女生举着烟花大笑，镜头贴近，蓝色烟花光在脸上跳动，0.8秒；
@图片2 电线杆下两人侧光剪影奔跑跳跃，仰拍，天空灰绿，0.8秒；
切换用帧闪白+轻微镜头晃动，节奏卡在音乐最强拍。禁止：任何文字字幕""",
    "2s", [f"img_{n:02d}" for n in [35, 33]])

# ========== OUTRO 消散 ==========
add_section_divider(doc, "═══  OUTRO 消散  |  46s – 50s  |  余韵·漂浮·归零  ═══", bg_color='0a1a2e', text_color='88ccff')

add_shot(doc, 10, "0:46 – 0:50", img(40), "特写·叠化·消散",
    "回到开场的少年戴耳机画面，缓慢叠化，画面逐渐过曝变白，如同记忆消散。最后只剩蓝天。",
    "@图片1 作为首帧和尾帧，少年戴耳机侧颜慢慢叠化，Slow push in极缓推向眼睛，画面从胶片质感逐渐过曝变为纯净蓝色，最后只剩蓝色虚化背景，镜头慢慢拉远消散，仿佛记忆溶解，最后一帧渐入蓝色纯色。4秒，16:9，音效：音乐渐弱，环境音消失。禁止：任何文字字幕",
    "4s", 'e8f8ff')

# ========== 附录：额外备用镜头 ==========
doc.add_paragraph()
add_section_divider(doc, "═══  备用镜头库（可替换主流程）  ═══", bg_color='2a2a2a', text_color='cccccc')

add_heading_style(doc, "以下镜头可根据节奏调整插入主流程", level=3, color='555555')

backup_shots = [
    (101, "备用A", img(2), "路边等待", "少女路边等车，Handheld跟拍",
     "@图片1 作为首帧，少女站在路边，Handheld手持跟拍从侧面，眼神飘向远方，色调温暖过曝，背景行人虚化流动。5秒，16:9。"),
    (102, "备用B", img(4), "自行车奔跑", "少年骑车快速穿越，跟拍",
     "@图片1 作为首帧，少年骑自行车，低角度侧面Tracking跟随，速度模糊，阳光在轮毂上碎裂。5秒，16:9。"),
    (103, "备用C", img(9), "铁路独站", "少年铁路边独立，建立镜头",
     "@图片1 作为首帧，少年站在铁路旁，固定全景镜头，画面构图对称，铁路延伸向远方消失点，色调冷蓝。5秒，16:9。"),
    (104, "备用D", img(31), "教室告白", "少年少女教室对视，两人之间光",
     "@图片1 作为首帧，空教室内少年少女面对面，窗外树影斑驳，慢推向两人之间的光，胶片质感蓝绿色调。5秒，16:9。"),
    (105, "备用E", img(36), "火车窗边少女", "少女坐火车发呆，窗外流动",
     "@图片1 作为首帧，少女撑脸坐火车，镜头贴近侧颜，窗外风景模糊向后流动，暖光在脸上忽明忽暗，伤感安静。5秒，16:9。"),
    (106, "备用F", img(26), "车内少年自拍", "两少年车内自拍，逆光",
     "@图片1 作为首帧，两少年在车内，强逆光从前挡风玻璃射入，手持相机感，鱼眼微畸变，胶片粒子。3秒，16:9。"),
    (107, "备用G", img(41), "教室混乱玩闹", "课堂混乱嬉闹，活力",
     "@图片1 作为首帧，教室里同学嬉闹，Handheld手持拍摄，镜头晃动，绿色调，窗外过曝，快切。3秒，16:9。"),
    (108, "备用H", img(8), "女生窗边", "女生窗边凝视，若有所思",
     "@图片1 作为首帧，女生站在窗边，侧光打亮半边脸，推镜头慢推，蓝色调，树影在玻璃上投影。5秒，16:9。"),
    (109, "备用I", img(21), "课堂发呆", "课堂上少年发呆凝视前方",
     "@图片1 作为首帧，教室少年发呆，鱼眼广角，同学们模糊，他清晰居中，慢推面部特写，窗外白光。3秒，16:9。"),
    (110, "备用J", img(15), "路灯夜景", "路灯下夜景，孤独感",
     "路灯下夜景，固定镜头，光晕在镜头上漫开，色调蓝黑，胶片粒子，孤独感，缓慢Zoom in路灯。3秒，16:9。"),
]

for sn, st, sp, desc, shot_type, prompt in backup_shots:
    add_shot(doc, sn, "备用", sp, shot_type, desc, prompt, "3-5s", 'f5f5f5')

# ========== 生产说明 ==========
doc.add_paragraph()
add_section_divider(doc, "═══  Seedance 生产说明  ═══", bg_color='1a1a2e', text_color='aaaaff')

notes = [
    "1. 所有提示词已按 Seedance 2.0 格式规范编写，可直接粘贴到即梦平台使用",
    "2. 组合分镜（三组）将多张图片一次性上传，参考@图片N引用语法，适合生成8-15秒连贯片段",
    "3. 单个分镜建议生成5秒视频，组合片段最长15秒（平台上限）",
    "4. 超过15秒的组合可使用「将@视频1延长Xs」方式续接",
    "5. 图片已按顺序重命名为img_01.png至img_41.png，存放在 mv_images 文件夹",
    "6. 建议按INTRO→BREAK→VERSE1→VERSE2→CHORUS→OUTRO顺序逐段生成再剪辑",
    "7. 总片长约50秒，由约10个主分镜+3个组合分镜构成",
]

for note in notes:
    p_note = doc.add_paragraph()
    r_note = p_note.add_run(note)
    r_note.font.size = Pt(10)
    r_note.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

doc.add_paragraph()

# 保存
doc.save(OUTPUT_PATH)
print(f"文档已保存: {OUTPUT_PATH}")
