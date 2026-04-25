# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面边距 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── 助手函数 ──
def set_font(run, name_cn='SimSun', name_en='Arial', size=11, bold=False, color=None, italic=False):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(text, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0,0,0)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_heading(text, size=14, bold=True, color=(30,30,30), space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_scene_header(text):
    """场景标题：深蓝底纹、白字"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(255,255,255))
    # 段落底纹
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2B4A7A')
    pPr.append(shd)
    return p

def add_shot(number, desc_cn, time_hint=''):
    """分镜描述行"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    label = f'【镜头 {number}】'
    if time_hint:
        label += f'  {time_hint}'
    run_label = p.add_run(label)
    set_font(run_label, size=11, bold=True, color=(180,80,0))
    run_desc = p.add_run('  ' + desc_cn)
    set_font(run_desc, size=11)
    return p

def add_dialogue(char_cn, char_jp, line_jp, line_cn, tone=''):
    """台词行（日文 + 中文对照）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # 角色名
    run_char = p.add_run(f'{char_cn}（{char_jp}）')
    set_font(run_char, size=10.5, bold=True, color=(100,0,120))
    if tone:
        run_tone = p.add_run(f'  [{tone}]')
        set_font(run_tone, size=9.5, italic=True, color=(120,120,120))
    p.add_run('\n')
    # 日文
    run_jp = p.add_run(f'  「{line_jp}」')
    set_font(run_jp, 'MS Gothic', 'Arial', size=11, bold=False)
    p.add_run('\n')
    # 中文对照
    run_cn = p.add_run(f'  （{line_cn}）')
    set_font(run_cn, size=10, italic=True, color=(80,80,80))
    return p

def add_action(text):
    """动作/画面说明"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f'▶ {text}')
    set_font(run, size=10.5, color=(50,50,50), italic=True)
    return p

def add_sfx(text):
    """音效说明"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'♪ 音效：{text}')
    set_font(run, size=10, color=(0,100,80))
    return p

def add_separator():
    p = doc.add_paragraph('─' * 55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(180,180,180)
    run.font.size = Pt(9)

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f'※ {text}')
    set_font(run, size=9.5, color=(130,130,130), italic=True)

# ══════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

add_title('渣男・番长・眼镜书虫', size=26, color=(30,60,120))
add_title('ダメ男・番長・眼鏡書虫', size=18, color=(80,80,80))

doc.add_paragraph()
add_title('剧　本　草　稿', size=16, color=(100,100,100))
doc.add_paragraph()
add_title('（喜剧短剧 · 日系3D动画风格）', size=12, color=(140,140,140))

for _ in range(2):
    doc.add_paragraph()

add_title('剧情概述', size=13, bold=True, color=(50,50,50))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '女主放学路上遭遇渣男借钱加羞辱，一怒之下暴揍一顿后进校。\n'
    '走廊偶遇学校"透明人"——读书眼镜男被混混欺负，钱包被抢。\n'
    '女主冷眼一扫将混混全部吓跑，转身淡定离开，什么也没说。\n'
    '眼镜男愣在原地，第一次被人帮助，低头自卑片刻后——\n'
    '鼓起勇气追上女主，故事的齿轮开始转动。'
)
set_font(run, size=11, color=(60,60,60))

doc.add_page_break()

# ══════════════════════════════════════════════
#  登场人物
# ══════════════════════════════════════════════
add_heading('【登场人物】', size=15, color=(20,20,80))

chars = [
    ('女主 · 神崎リオ（かんざき りお）', '高中三年生。暗红色短发，绿色大耳环，黑色颈圈，粉色V领针织衫配百褶裙。\n外表强势、情绪外露，骨子里对真心有所期待。口头禅是各种海鲜系骂人词。'),
    ('渣男 · 黑田タクミ（くろだ たくみ）', '寸头，银耳钉，藏青西装外套+红色T恤。\n油嘴滑舌、自恋，金钱与肉欲挂帅，但有一分贱到底的真实。'),
    ('眼镜男 · 浅野ケン（あさの けん）', '黑色短发呆毛，圆框眼镜，黑色学生制服。\n安静书虫型，班里透明人，内心敏感。被欺负从不还手，只会低头。'),
    ('混混三人组', '学校问题学生，总是集体行动，欺负眼镜男取乐。遇强则弱。'),
]

for name, desc in chars:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    run_name = p.add_run(f'◆ {name}\n')
    set_font(run_name, size=11, bold=True, color=(30,60,120))
    run_desc = p.add_run(f'   {desc}')
    set_font(run_desc, size=10.5)

doc.add_page_break()

# ══════════════════════════════════════════════
#  ACT 1 · 校门口渣男遭遇战
# ══════════════════════════════════════════════
add_heading('ACT 1   校门口渣男遭遇战', size=15, color=(150,30,0))
add_note('时长参考：约 45 秒 / 分三段生成  |  场景：学校外侧砖墙街道，蓝天晴天')

# ── SCENE 1-1 ──
add_scene_header('SCENE 1-1   借钱谈判·女主炸毛')

add_shot(1, '女主面部极近景正对镜头，背景蓝天，眉毛紧皱、双眼瞪大，绿色耳环微晃。', '0–2s')
add_dialogue('女主','リオ','はぁ？','啥？','愤怒质问')

add_shot(2, '过肩镜头——从渣男背后拍，越过其肩膀看向对面女主，女主冷淡带刺，眼神直盯对方。', '2–6s')
add_dialogue('女主','リオ','だから、金貸してくれねぇなら——','所以说，你不借我钱的话——','冷淡强硬')
add_dialogue('女主','リオ','今日のデートはなし！','今天约会就算了！','斩钉截铁')

add_shot(3, '反打切换，渣男靠着砖墙，侧身斜眼，嘴角微勾，一副漫不经心的轻浮表情。', '7–11s')
add_dialogue('渣男','タクミ','でも、いい加減やらしてくれるならいいぜ。','但你要是差不多愿意让我上了也行。','轻佻冷淡')

add_shot(4, '渣男继续扬起下巴，眯眼，伸手指着女主方向，充满挑衅。', '11–14s')
add_dialogue('渣男','タクミ','ダブルホテル代は、そっち持ちで。','双人房费用你来出。','意味深长')

add_shot(5, '女主瞳孔震动极近特写，怒火上涌，下一秒攥拳——出击。', '14–15s')
add_action('女主出拳，漫画冲击波特效瞬间爆发。')
add_sfx('弦乐短促升起 → 重击音')

add_separator()

# ── SCENE 1-2 ──
add_scene_header('SCENE 1-2   漫画式街头暴打')

add_shot(6, '女主振臂冲向渣男，拳头特写，漫画冲击线爆开。', '0–3s')
add_sfx('ガッチ！× 2')

add_shot(7, '渣男被打到踉跄靠墙，女主连续出拳；俯拍全景，两人激烈拉扯，快切。', '3–8s')
add_sfx('连续 ガッチ！节奏打击音 × 3')

add_shot(8, '短暂距离拉开，女主喘气叉腰，怒气值拉满；渣男扶墙整理衣领，嘴角犹带轻浮笑意。', '8–13s')
add_action('渣男整理西装，抬头，语气嘲讽中带冷淡。')
add_dialogue('渣男','タクミ','もういいわ。おめぇなんかよ——','我不干了。你这种人——','毒舌冷淡')
add_dialogue('渣男','タクミ','金もない、やらせてもくれない女！興味ねぇわ！','没钱还不给上的女人！对你没兴趣！','嘲讽离场')

add_shot(9, '渣男转身离开，插兜，不回头。女主站在原地，气势汹汹指着其背影。', '13–15s')
add_action('女主连环骂人收尾，每句一个硬切。')
add_dialogue('女主','リオ','バカ！　タコ！　イカ！　マグロ！',
            '笨蛋！章鱼！乌贼！金枪鱼！','连环怒骂')
add_sfx('每次硬切配漫画爆破短音；结尾一片樱花飘过，街道安静。')

doc.add_page_break()

# ══════════════════════════════════════════════
#  ACT 2 · 进校·走廊独白
# ══════════════════════════════════════════════
add_heading('ACT 2   进校·走廊独白', size=15, color=(150,30,0))
add_note('时长参考：约 10 秒  |  场景：日式学校正门 → 走廊')

add_scene_header('SCENE 2-1   烦闷走廊')

add_shot(10, '女主进入校门，鞋子特写，步伐沉重，节奏缓慢，地板回响。', '0–3s')

add_shot(11, '走廊中景，女主微低着头，表情气鼓鼓，独自沿走廊行走，光线从窗户射入。', '3–7s')
add_dialogue('女主','リオ（内心独白）',
            'うちの人生では、もうまともな男には出会えないよ……',
            '我这辈子，大概再也遇不到一个正经男人了……','气鼓鼓内心独白')

add_shot(12, '女主自言自语，一边走一边摆手，表情从气恼变成无奈。', '7–10s')
add_dialogue('女主','リオ',
            'ていうかな……嘘だろ、誰かよ。せちがらい世の中だ。',
            '什么嘛……不是吧，有没有人啊。真是个冷漠的世道。','无奈吐槽')
add_action('走廊另一头传来喧嚣声，女主脚步停住，侧耳。')
add_sfx('远处吵闹声混入，女主停步。')

doc.add_page_break()

# ══════════════════════════════════════════════
#  ACT 3 · 发现欺凌·冷眼震慑
# ══════════════════════════════════════════════
add_heading('ACT 3   发现欺凌·冷眼震慑', size=15, color=(150,30,0))
add_note('时长参考：约 25 秒  |  场景：学校教室/走廊侧门')

add_scene_header('SCENE 3-1   混混欺负眼镜男')

add_shot(13, '教室内，三个混混将眼镜男团团围住。混混A把眼镜男书桌上的书推到地上，混混B把玩从其口袋里抢走的钱包。', '0–4s')
add_dialogue('混混A','チンピラA',
            '中に磁石入れてみ。',
            '（拍打眼镜男肩膀）试试往里面放磁铁。','挑衅取乐')
add_dialogue('混混B','チンピラB',
            '痛いっしょ、これ。',
            '（捏着眼镜男手腕）这很痛吧？','嘲笑')

add_shot(14, '眼镜男低头，眼镜上还摊开着一本杂志，表情僵硬，不敢还击。', '4–7s')
add_dialogue('混混C','チンピラC',
            'なに読んでんの？それ。',
            '你在看什么啊？这个。','嘲讽翻书')
add_dialogue('眼镜男','ケン',
            'え？　あ、あの……',
            '啊？那个……','紧张结巴')

add_shot(15, '混混A将眼镜男的钱包高高举起，准备朝走廊扔出去。', '7–10s')
add_action('钱包飞出画面。眼镜男本能起身想要追，被混混B摁回椅子上。')

add_separator()

add_scene_header('SCENE 3-2   女主登场·一眼定乾坤')

add_shot(16, '教室门口——女主出现在侧门阴影处，半个身影，只能看到她的眼神。\n从教室内往外看的视角，逆光，她的眼睛在光线中清晰可见。', '0–2s')
add_action('全场安静。混混三人组意识到有人，回头。')

add_shot(17, '正面特写：女主冷漠走入室内，视线从左扫向右，平静而具有压迫感，没有任何威胁性语言，只是看着。\n她的绿色耳环在光线下微微晃动。', '2–6s')
add_action('混混三人与女主眼神相接，愣住。')

add_shot(18, '混混A、B、C三人对视，无声交流，感受到压力，慢慢向后退步。', '6–9s')
add_action('混混C把钱包往地上一扔，三人转身快步离开，嘴里嘟囔。')
add_dialogue('混混A','チンピラA',
            'ちっ……行こうぜ。',
            '（吸牙）……走吧。','被震慑夹尾溜走')

add_shot(19, '女主目送三人离去，视线重新回到教室，扫了一眼眼镜男，表情没什么变化，\n她从地上捡起滑落的钱包，放到眼镜男桌上，不说话。', '9–13s')
add_action('女主转身，走向门口，背对眼镜男。')
add_dialogue('女主','リオ',
            '静かにしかいねぇのか、この世は。',
            '这世上，就没有个消停的人吗。','对空气吐槽，非对眼镜男说')

add_shot(20, '女主走出教室，背影消失在走廊拐角。教室只剩眼镜男一人。', '13–15s')
add_sfx('脚步声消失，教室安静，窗外风声与远处树叶沙沙声。')

doc.add_page_break()

# ══════════════════════════════════════════════
#  ACT 4 · 自卑·鼓起勇气
# ══════════════════════════════════════════════
add_heading('ACT 4   愣神·自卑·鼓起勇气', size=15, color=(150,30,0))
add_note('时长参考：约 15 秒  |  场景：教室 → 走廊')

add_scene_header('SCENE 4-1   眼镜男愣神')

add_shot(21, '眼镜男坐在原地，低头看着桌上自己的钱包，一动不动。\n窗外蓝天，光线在他身上静止。', '0–4s')
add_action('慢镜头感，眼镜男垂下眼帘，嘴唇微动，没有声音。')
add_dialogue('眼镜男（内心独白）','ケン（心の声）',
            '……なんで。',
            '……为什么。','茫然、不可置信地被帮助')

add_shot(22, '眼镜男侧面近景，从窗边看去，光打在他眼镜镜片上，映出走廊方向的光亮。\n他慢慢抬起头，看向门口方向。', '4–8s')
add_dialogue('眼镜男（内心独白）','ケン（心の声）',
            '誰も……助けてくれたことなんてなかったのに。',
            '从来没有人……帮过我的。','第一次被帮助的震动')

add_separator()

add_scene_header('SCENE 4-2   追上女主')

add_shot(23, '眼镜男站起来，椅子摩擦地面的声音划破寂静，他快步走出教室。', '0–2s')
add_sfx('椅子摩擦声 + 奔跑脚步声')

add_shot(24, '走廊，眼镜男加速小跑，眼镜微微晃动，表情紧绷而下定决心。', '2–5s')
add_action('前方走廊拐角，女主的背影刚好出现在视野边缘。')

add_shot(25, '眼镜男追上，在女主身后几步停下，喘了口气，鼓起勇气开口。\n女主没有回头，继续走。', '5–9s')
add_dialogue('眼镜男','ケン',
            'あ、あの……っ！',
            '那、那个……！','气喘吁吁，鼓足勇气')

add_shot(26, '女主放慢脚步，没有停，也没有回头，侧脸在光线里若有若无。', '9–12s')
add_action('女主没有说话，只是步伐稍微慢了下来——这已经是她的回应。')

add_shot(27, '眼镜男深吸一口气，握紧拳头，把想说的话挤出来。', '12–15s')
add_dialogue('眼镜男','ケン',
            'さっきは……ありがとう、ございました。',
            '刚才……谢、谢谢你。','诚恳、略显笨拙的道谢')
add_action('女主停住。背影静止在走廊光线里，没有回头。')
add_sfx('走廊安静，远处教室噪音渐淡，风声从走廊窗户吹过，头发轻扬。')
add_note('此处为全片结尾定格——女主是否回头，留白给观众想象。')

doc.add_page_break()

# ══════════════════════════════════════════════
#  台词完整对照表
# ══════════════════════════════════════════════
add_heading('台词完整对照表', size=15, color=(20,20,80))
add_note('★ 标注行为从视频原声提取，其余为剧本创作。')

# 表格
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '场景'
hdr[1].text = '角色'
hdr[2].text = '日文台词'
hdr[3].text = '中文对照'
hdr[4].text = '备注'

for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)

dialogues = [
    ('ACT1-1','女主','はぁ？','啥？','★ 原声'),
    ('ACT1-1','女主','だから、金貸してくれねぇなら、今日のデートはなし！','所以你不借我钱的话，今天约会取消！','★ 原声'),
    ('ACT1-1','渣男','でも、いい加減やらしてくれるならいいぜ。','但你要是差不多愿意让我上了也行。','★ 原声'),
    ('ACT1-1','渣男','ダブルホテル代は、そっち持ちで。','双人房费用你来出。','★ 原声'),
    ('ACT1-2','渣男','もういいわ。おめぇなんかよ——金もない、やらせてもくれない女！興味ねぇわ！','我不干了。你这种人——没钱还不给上的女人！对你没兴趣！','★ 原声'),
    ('ACT1-2','女主','バカ！タコ！イカ！マグロ！','笨蛋！章鱼！乌贼！金枪鱼！','★ 原声'),
    ('ACT2','女主（独白）','うちの人生では、もうまともな男には出会えないよ……','我这辈子，大概再也遇不到正经男人了……','参考原声改写'),
    ('ACT2','女主','ていうかな……嘘だろ、誰かよ。せちがらい世の中だ。','什么嘛……不是吧。真是个冷漠的世道。','★ 参考原声'),
    ('ACT3','混混A','中に磁石入れてみ。','试试往里面放磁铁。','★ 原声'),
    ('ACT3','混混B','痛いっしょ、これ。','这很痛吧？','★ 原声'),
    ('ACT3','混混C','なに読んでんの？それ。','你在看什么啊？','★ 原声'),
    ('ACT3','眼镜男','え？あ、あの……','啊？那个……','★ 原声'),
    ('ACT3','混混A','ちっ……行こうぜ。','（吸牙）……走吧。','剧本创作'),
    ('ACT3','女主','静かにしかいねぇのか、この世は。','这世上，就没有个消停的人吗。','参考原声改写'),
    ('ACT4','眼镜男（内心）','……なんで。','……为什么。','剧本创作'),
    ('ACT4','眼镜男（内心）','誰も……助けてくれたことなんてなかったのに。','从来没有人帮过我的。','剧本创作'),
    ('ACT4','眼镜男','あ、あの……っ！','那、那个……！','剧本创作'),
    ('ACT4','眼镜男','さっきは……ありがとう、ございました。','刚才……谢、谢谢你。','剧本创作'),
]

for scene, char, jp, cn, note in dialogues:
    row = table.add_row().cells
    row[0].text = scene
    row[1].text = char
    row[2].text = jp
    row[3].text = cn
    row[4].text = note
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)

# ══════════════════════════════════════════════
#  创作说明
# ══════════════════════════════════════════════
doc.add_page_break()
add_heading('创作说明 / 制作备注', size=14, color=(20,20,80))

notes_list = [
    '本剧本基于视频《4月18日-1.mp4》及《4月18日(1)-1.mp4》进行分析与创作，原声日文台词使用 faster-whisper small 模型提取（日语置信度 100%）。',
    '男主"浅野ケン"与女主"神崎リオ"为剧本虚构命名，可根据制作需求修改。',
    'ACT1 三个场景对应已生成的 Seedance 提示词，可直接使用@图片1（女主）/@图片2（渣男）调用。',
    'ACT3 混混角色建议使用 Seedance 随机生成（不锁定），以避免角色过多锁定影响生成质量。',
    '眼镜男（浅野ケン）角色造型参考视频中书虫形象，建议另行准备参考图锁定。',
    'ACT4 结尾定格处理参考岩井俊二式留白风格，女主是否回头由制作方决定。',
    '全剧本台词均使用关西/江户混合口语体，具口语化与年龄感，适配3D日系动画气质。',
]

for i, note in enumerate(notes_list, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    run = p.add_run(f'{i}. {note}')
    set_font(run, size=10.5)

# ══════════════════════════════════════════════
#  保存
# ══════════════════════════════════════════════
out_path = 'E:/picture2/渣男番长眼镜书虫_剧本.docx'
doc.save(out_path)
print(f'Done: {out_path}')
